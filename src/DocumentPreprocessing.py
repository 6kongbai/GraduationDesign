from types import NoneType
from sentence_transformers import SentenceTransformer
from docx.text.paragraph import Paragraph
from docx.parts.image import ImagePart
from docx.table import _Cell, Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from src.CodeType import CodeType
from src.MyParagraph import MyParagraph
import torch
import pandas as pd
import docx
import os

device = torch.device("cuda")
# 加载预训练模型并将其移动到 GPU
model = SentenceTransformer('DMetaSoul/Dmeta-embedding').to(device)


class MyDocument:
    """
    文档
    """

    def __init__(self, path: str, index: int, code_type: str):
        """
        初始化文档
        :param path:文档路径
        :param index:文档序号
        :param code_type:代码类型 0:java 1:python 2:c 3:c++
        """
        # 文档路径
        self.path: str = path
        # 文档索引
        self.doc_index: int = index
        # 文档对象
        self.doc: docx.Document = docx.Document(path)
        # 代码类型
        self.code_type: CodeType = CodeType.get_code_type(code_type)
        # 文档段落
        self.paragraphs: pd.DataFrame = pd.DataFrame()
        # 文档句子嵌入向量
        self.embeddings: torch.Tensor = torch.Tensor()
        # 文档标题
        self.title: str = os.path.basename(path)
        # 文档中的源代码
        self.codes: list = []
        # 解析文档
        self._parse_word()

    def paser_part(self, block):
        """
        处理文档部分，根据输入的块类型（段落或表格），返回相应的处理结果。

        :param block: 需要处理的文档部分，可以是段落（CT_P）或者表格（CT_Tbl）。
        :return: 返回处理后的对象，如果是图片段落则返回ImagePart对象，如果是普通段落则返回Paragraph对象，
                 如果是表格则返回Table对象。
        """
        if isinstance(block, CT_P):
            # 处理段落类型的块
            para = Paragraph(block, self.doc)
            if is_image(para, self.doc):
                # 如果段落中包含图片，则返回图片部分对象
                return get_ImagePart(para, self.doc)
            return Paragraph(block, self.doc)
        elif isinstance(block, CT_Tbl):
            # 处理表格类型的块
            return Table(block, self.doc)

    def block_process(self, block, block_index, rows_to_append) -> CodeType:
        """
        处理文档中的一个区块，识别并处理该区块的特定内容。

        :param block: 需要处理的文档区块，可以是段落、图片或表格等。
        :param block_index: 区块在文档中的索引。
        :param rows_to_append: 用于收集需要追加到结果中的行数据的列表。
        :return: 返回识别到的代码类型，如果没有识别到则返回NULL。
        """
        if isinstance(block, Paragraph):
            # 判断当前区块是否为正文段落
            if block.style.name == 'Normal':
                paragraph = MyParagraph(block.text, self.code_type)
                flag = paragraph.check()
                # 如果段落中包含需要的代码类型且长度超过16个字符，则将其句子信息追加到rows_to_append中
                if flag != CodeType.NULL and flag != self.code_type and len(paragraph.text) > 16:
                    rows_to_append.extend([
                        {
                            'document_index': self.doc_index,
                            'paragraph_index': block_index,
                            'sent_index': sent_index,
                            'sent': sent
                        }
                        for sent_index, sent in enumerate(paragraph.sentences)
                    ])
                return flag

            # 判断是否为标题。当前只处理Heading 1，后续可以根据需要扩展处理其他级别的标题
            elif block.style.name == 'Heading 1':
                pass
        elif isinstance(block, ImagePart):
            # 处理图片区块的逻辑预留，当前未实现
            pass
        elif isinstance(block, Table):
            # 处理表格区块的逻辑预留，当前未实现
            pass

    def code_process(self, block) -> CodeType:
        """
        处理代码段。
        对输入的文本块进行处理，判断其语言类型，并将其添加到当前的代码片段列表中。如果该段代码不是中文，
        则将其追加到当前的代码片段中，并以换行符分隔。

        :param block: 包含待处理代码文本的块对象。
        :return CodeType: 检测到的代码语言类型。
        """
        paragraph = MyParagraph(block.text, self.code_type)  # 创建一个段落对象，用于分析和处理输入的代码块
        flag = paragraph.check()  # 检查代码段的语言类型
        if flag != CodeType.CHINESE:
            self.codes[-1] += paragraph.text + '\n'  # 如果不是中文，将代码段追加到当前代码片段列表的最后一个元素中
        return flag  # 返回检测到的代码语言类型

    def _parse_word(self):
        """
        解析word文档
        :return:
        """
        rows_to_append = []
        it = iter(self.doc.element.body)  # 创建一个迭代器
        block_index = 0
        flag = True  # 当为True时，表示当前部分为非代码部分，否则为代码部分
        while True:
            try:
                part = next(it)  # 获取下一个部分
                block = self.paser_part(part)  # 获取当前部分
                if isinstance(block, NoneType):
                    continue
                if flag:
                    part_type = self.block_process(block, block_index, rows_to_append)  # 处理当前部分
                    if part_type == self.code_type:
                        self.codes.append("""""")
                        self.code_process(block)
                        flag = False
                else:
                    part_type = self.code_process(block)
                    if part_type == CodeType.CHINESE:
                        flag = True
                block_index += 1
            except StopIteration:
                break  # 迭代结束时退出循环
        self.paragraphs = pd.DataFrame(rows_to_append, columns=['document_index', 'paragraph_index', 'sent_index', 'sent'])
        # 计算句子嵌入向量
        self.calculate_embeddings()

    def calculate_embeddings(self):
        """
        计算并存储每个句子的嵌入表示。

        该方法提取每个段落中的句子文本，并使用预训练的模型计算它们的嵌入向量，
        最后将这些嵌入向量存储起来以便后续使用。

        参数:
        self - 对象自身的引用。

        返回值:
        无
        """
        # 提取句子文本，保留段落号和句子号信息
        sentences = self.paragraphs['sent'].values
        # 使用model.encode计算句子嵌入矩阵，并转移到指定的计算设备上
        self.embeddings = model.encode(sentences, convert_to_tensor=True).to(device)

    def calculate_similarity(self, other_doc: 'MyDocument'):
        """
        计算当前文档与另一个文档之间的相似度。

        参数:
        - other_doc: 'MyDocument' 类型，要与当前文档进行相似度比较的另一个文档对象。

        说明:
        此方法首先计算两个文档嵌入向量之间的相似度矩阵，然后找出每个句子在另一个文档中最相似的句子，
        并打印出相似度高于0.85的句子对及其相似度分数。
        """
        # 计算两个文档的嵌入向量之间的相似度矩阵
        similarity_matrix = calculate(self.embeddings, other_doc.embeddings)
        # 在相似度矩阵的每一行中找到最大值及其索引，表示每个句子在另一个文档中的最佳匹配
        max_values, indices = torch.max(similarity_matrix, dim=1)
        # 遍历并打印出相似度高于0.85的句子对及其相似度分数
        for i, j in enumerate(indices):
            if max_values[i] > 0.85:
                print('原句子:', self.get_sentence_info(i)['sent'],
                      '\n相似句子:', other_doc.get_sentence_info(j.item())['sent'],
                      f'\n相似分数:{max_values[i]:.2f}\n')

    def get_sentence_info(self, index):
        return self.paragraphs.iloc[index, :]


# 该行只能有一个图片
def is_image(graph: Paragraph, doc: docx.Document):
    images = graph._element.xpath('.//pic:pic')  # 获取所有图片
    for image in images:
        for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
            part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
            if isinstance(part, ImagePart):
                return True
    return False


# 获取图片（该行只能有一个图片）
def get_ImagePart(graph: Paragraph, doc: docx.Document):
    images = graph._element.xpath('.//pic:pic')  # 获取所有图片
    for image in images:
        for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
            part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
            if isinstance(part, ImagePart):
                return part
    return None


def calculate(embeddings1, embeddings2):
    # 计算两两相似度
    return embeddings1 @ embeddings2.T


if __name__ == '__main__':
    name = 'B20200103213-唐睿智-21软件01-软件工程基础实训II(2023)-校内集中实训课程实训报告.docx'
    ROOT_DIR_P = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))  # 项目根目录
    word_path = os.path.join(ROOT_DIR_P, f"files/{name}")  # pdf文件路径及文件名
    d1 = MyDocument(word_path, 0, 'java')
    print(d1.paragraphs)
